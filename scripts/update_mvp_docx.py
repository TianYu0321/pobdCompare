from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


AI_DIR = Path(r"D:\ai")
DOCX_NAME_FRAGMENT = "PoE2 BD 差异比较与收益模拟工具 MVP 详细实现设计方案"
BACKUP_SUFFIX = ".pre-20260620-release-backup.docx"


def insert_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    paragraph._p.addnext(element)
    inserted = Paragraph(element, paragraph._parent)
    if style:
        inserted.style = style
    inserted.add_run(text)
    return inserted


def main() -> None:
    if len(sys.argv) > 1:
        target = Path(sys.argv[1])
    else:
        candidates = [
            path
            for path in AI_DIR.glob("*.docx")
            if DOCX_NAME_FRAGMENT in path.name and ".pre-" not in path.name
        ]
        if len(candidates) != 1:
            raise RuntimeError(f"Expected one target DOCX, found: {candidates}")
        target = candidates[0]

    backup = target.with_name(target.stem + BACKUP_SUFFIX)
    if not backup.exists():
        shutil.copy2(target, backup)

    document = Document(str(target))

    revision_heading_found = False
    for paragraph in document.paragraphs:
        if paragraph.text in {
            "2026-06-19 P3/MVP 范围修订（优先级高于原文）",
            "2026-06-20 P3/MVP 范围与实施状态修订（优先级高于原文）",
        }:
            paragraph.text = "2026-06-20 P3/MVP 范围与实施状态修订（优先级高于原文）"
            revision_heading_found = True
            break
    if not revision_heading_found:
        raise RuntimeError("Could not find the P3/MVP revision heading")

    marker = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if "后置到 P3.5/P4" in paragraph.text
        ),
        None,
    )
    if marker is None:
        raise RuntimeError("Could not find the P3/P4 deferral marker")

    if not any("2026-06-20 实施状态" in paragraph.text for paragraph in document.paragraphs):
        current = insert_after(marker, "2026-06-20 实施状态", "Heading 1")
        status_lines = [
            "1. 本地 Fastify API、SSE、双 BD/单 BD 工作台、固定装备槽位、装备详情抽屉、Variant revision、undo/redo/reset 与三类天赋收益榜已经落地。",
            "2. WeGame Profile API、版本化 MappingCatalog、精确映射 blocker、PoB2 原生 ImportTab、SaveDB/reload 和 baseline 验证链已经实现。",
            "3. poe.ninja 使用动态 snapshot 和 pathOfBuildingExport 导入，不从页面文本反推构筑。",
            "4. 真实 PoB2 XML 已完成浏览器端到端验证；构建通过，普通测试 381 项通过，原生 WeGame bridge 集成测试需显式开启 POB2_INTEGRATION=1。",
            "5. 发布前仍须保存一条当前真实 WeGame 链接和一条当前真实 poe.ninja 链接的端到端 calculable 验收证据。未通过精确映射或 PoB2 round-trip 的数据继续保持 normalized，绝不输出假收益。",
        ]
        for line in status_lines:
            current = insert_after(current, line)

    replacements = {
        "阶段 5：WeGame Adapter": "阶段 5：WeGame Adapter（P3/MVP 已实现，持续维护映射目录）",
        "完成该阶段后，再进入 WeGame Adapter 和 Agent 报告层。":
            "当前已进入完整 P3/MVP 工作台阶段；WeGame Adapter 属于 MVP 已实现范围，后续重点是持续维护精确映射目录并补齐真实链接验收证据。",
        "4. 真实 PoB2 XML 已完成浏览器端到端验证；构建通过，普通测试 381 项通过，原生 WeGame bridge 集成测试需显式开启 POB2_INTEGRATION=1。":
            "4. 真实 PoB2 XML 已完成浏览器端到端验证；构建通过，普通测试 381 项通过；另以 POB2_INTEGRATION=1 开启的原生 WeGame bridge 集成测试 1/1 通过。",
        "5. 发布前仍须保存一条当前真实 WeGame 链接和一条当前真实 poe.ninja 链接的端到端 calculable 验收证据。未通过精确映射或 PoB2 round-trip 的数据继续保持 normalized，绝不输出假收益。":
            "5. 2026-06-20 已用当前真实 WeGame 分享角色完成 calculable、PoB2 round-trip、双 BD 对比和三类天赋榜验收；装备 16/16、技能条目 53/53、词条 99/99、天赋 104/104 精确映射。仍需补充当前真实 poe.ninja 链接验收；未通过精确映射或 PoB2 round-trip 的数据继续保持 normalized，绝不输出假收益。",
    }
    for paragraph in document.paragraphs:
        replacement = replacements.get(paragraph.text)
        if replacement:
            paragraph.text = replacement

    document.save(str(target))
    print(target)
    print(backup)


if __name__ == "__main__":
    main()
